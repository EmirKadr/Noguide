import os
import gradio as gr
from docx import Document
import fitz
import base64
import re
import html  # <-- NYTT för att visa kod korrekt
from io import BytesIO

# === Befintliga PDF/DOCX funktioner ===

def extract_text_from_pdf(path):
    try:
        with fitz.open(path) as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception as e:
        return f"PDF ERROR: {e}"

def extract_text_from_docx(path):
    try:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"DOCX ERROR: {e}"

def load_documents(folder):
    docs = []
    for filename in os.listdir(folder):
        path = os.path.join(folder, filename)
        if filename.lower().endswith(".pdf"):
            content = extract_text_from_pdf(path)
        elif filename.lower().endswith(".docx"):
            content = extract_text_from_docx(path)
        else:
            continue
        docs.append({"filename": filename, "content": content, "path": path})
    return docs

documents = load_documents("docs")

# === Helper: extract snippet ===

def extract_context_snippet(text, query, max_chars=600):
    text = text.replace("\n", " ")
    match = re.search(re.escape(query), text, flags=re.IGNORECASE)
    if not match:
        return None

    start = max(match.start() - max_chars // 2, 0)
    end = min(match.end() + max_chars // 2, len(text))
    snippet = text[start:end]

    if start > 0:
        snippet = "…" + snippet
    if end < len(text):
        snippet += "…"

    def highlight(match):
        return f"<mark>{match.group(0)}</mark>"
    return re.sub(f"({re.escape(query)})", highlight, snippet, flags=re.IGNORECASE).strip()

# === search_documents funktion ===

def search_documents(query, visible_count=5):
    if not query or len(query.strip()) < 2:
        return "❗️ Skriv minst 2 tecken för att söka.", gr.update(visible=False)

    query = query.strip()
    results = []

    for doc in documents:
        score = 0
        filename_match = query.lower() in doc['filename'].lower()
        content_match = query.lower() in doc['content'].lower()

        if filename_match:
            score += 1000
        if content_match:
            score += 10

        if score > 0:
            results.append((doc, score, filename_match))

    results.sort(key=lambda x: (x[1], x[2]), reverse=True)

    html_output = ""
    shown = 0
    for doc, score, filename_match in results:
        if shown >= visible_count:
            break

        highlighted_filename = re.sub(
            f"({re.escape(query)})",
            r"<mark>\1</mark>",
            doc['filename'],
            flags=re.IGNORECASE
        )

        snippet = extract_context_snippet(doc["content"], query)

        if not snippet and filename_match:
            snippet = f"<div style='color:green'><b>Sökordet hittades i filnamnet.</b></div>"

        html_output += f"<h4>📄 {highlighted_filename}</h4>"
        if snippet:
            html_output += f"<div style='background-color:#f6f6f6;padding:10px;border-radius:5px;margin-bottom:5px;'>{snippet}</div>"
        else:
            html_output += f"<p style='color:gray;'>⚠️ Ingen tydlig träfftext hittades.</p>"

        html_output += f"<p>🔍 <b>Matchningspoäng:</b> {round(score, 1)} "

        try:
            with open(doc["path"], "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
                html_output += f"📥 <a href='data:application/octet-stream;base64,{b64}' download='{doc['filename']}'>Ladda ner filen</a></p><hr>"
        except:
            html_output += "❌ Gick inte att ladda filen</p><hr>"

        shown += 1

    show_more_visible = shown < len(results)
    return html_output if html_output else "❌ Inga träffar hittades.", gr.update(visible=show_more_visible)

# === NY: Ladda och strukturera Word-dokument med rubrik + text + bilder ===

def parse_word_sections(path):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    doc = Document(path)
    sections = []
    current_section = {"heading": "", "text": "", "images": []}

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current_section["heading"]:
                sections.append(current_section)
            current_section = {"heading": para.text, "text": "", "images": []}
        else:
            current_section["text"] += para.text + "\n"

        # Kolla om para har runs med bilder
        for run in para.runs:
            if run.element.xpath('.//a:blip'):
                # Hämta relation id
                blip = run.element.xpath('.//a:blip')[0]
                rEmbed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                image_part = doc.part.related_parts[rEmbed]
                image_bytes = image_part.blob
                img_b64 = base64.b64encode(image_bytes).decode()
                current_section["images"].append(img_b64)

    if current_section["heading"]:
        sections.append(current_section)

    return sections


word_sections = parse_word_sections(os.path.join("quickSearch", "Bibliotek.docx"))

# === NY: Sökfunktion för Word-dokument ===

def search_word_doc(query, visible_count=5):
    if not query or len(query.strip()) < 2:
        return "❗️ Skriv minst 2 tecken för att söka."

    query = query.strip()
    results = []

    for section in word_sections:
        score = 0
        heading_match = query.lower() in section['heading'].lower()
        text_match = query.lower() in section['text'].lower()

        if heading_match:
            score += 60
        if text_match:
            score += 10

        if score > 0:
            results.append((section, score, heading_match))

    results.sort(key=lambda x: (x[1], x[2]), reverse=True)

    html_output = ""
    shown = 0
    for section, score, heading_match in results:
        if shown >= visible_count:
            break

        highlighted_heading = re.sub(
            f"({re.escape(query)})",
            r"<mark>\1</mark>",
            section['heading'],
            flags=re.IGNORECASE
        )

        snippet = extract_context_snippet(section["text"], query)
        if not snippet and heading_match:
            snippet = f"<div style='color:green'><b>Sökordet hittades i rubriken.</b></div>"

        # Bilder i base64
        images_html = ""
        for img_b64 in section['images']:
            images_html += f"<img src='data:image/png;base64,{img_b64}' style='max-width:100%;'><br>"

        # NYTT: Visa hela texten med <pre> för kodformat
        escaped_text = html.escape(section['text'])
        html_output += f"<h4>📑 {highlighted_heading}</h4>"
        html_output += f"<div style='background-color:#f6f6f6;padding:10px;border-radius:5px;margin-bottom:5px;'>{snippet}</div>"
        html_output += f"<details><summary>▶️ Läs mer</summary><pre style='white-space: pre-wrap;'>{escaped_text}</pre>{images_html}</details>"
        html_output += f"<p>🔍 <b>Matchningspoäng:</b> {round(score, 1)}</p><hr>"

        shown += 1

    return html_output if html_output else "❌ Inga träffar hittades."

# === Gradio UI ===

with gr.Blocks() as demo:
    gr.Markdown("# 📚 NoWaste Dokumentbibliotek")

    with gr.Tab("Sök i dokumentmapp"):
        query1 = gr.Textbox(label="🔍 Sök i dokumentmapp", placeholder="Ex: inventering, pall, artikelnummer")
        output1 = gr.HTML()
        visible_count1 = gr.State(5)
        show_more_btn1 = gr.Button("⬇️ Visa fler", visible=False)

        def show_more_results(query, visible_count):
            return search_documents(query, visible_count + 5) + (visible_count + 5,)

        query1.change(fn=search_documents, inputs=[query1, visible_count1], outputs=[output1, show_more_btn1])
        show_more_btn1.click(fn=show_more_results, inputs=[query1, visible_count1], outputs=[output1, show_more_btn1, visible_count1])

    with gr.Tab("Sök i Bibliotek.docx"):
        query2 = gr.Textbox(label="🔍 Sök i Bibliotek.docx", placeholder="Ex: lager, capabilities, brandfarligt")
        output2 = gr.HTML()

        query2.change(fn=search_word_doc, inputs=query2, outputs=output2)

if __name__ == "__main__":
    demo.launch()
