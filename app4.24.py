import os
import re
import base64
import html

import gradio as gr

from docx import Document
import fitz  # PyMuPDF

from rapidfuzz import fuzz
from pygments import highlight
from pygments import lexers
from pygments.lexers import PythonLexer, SqlLexer, JsonLexer, TextLexer
from pygments.formatters import HtmlFormatter

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# =========================
#    PDF/DOCX – laddning
# =========================
def extract_text_from_pdf(path):
    try:
        with fitz.open(path) as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception as e:
        return f"PDF ERROR: {e}"

def extract_text_from_docx(path):
    try:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"DOCX ERROR: {e}"

def load_documents(folder):
    docs = []
    if not os.path.exists(folder):
        return docs
    for filename in os.listdir(folder):
        path = os.path.join(folder, filename)
        if filename.lower().endswith(".pdf"):
            content = extract_text_from_pdf(path)
        elif filename.lower().endswith(".docx"):
            content = extract_text_from_docx(path)
        else:
            continue
        docs.append({"filename": filename, "content": content, "path": path})
    return docs

documents = load_documents(os.path.join(BASE_DIR, "docs"))

# =========================
#   Hjälpare: snippets
# =========================
def extract_context_snippet(text, query, max_chars=600):
    text = text.replace("\n", " ")
    match = re.search(re.escape(query), text, flags=re.IGNORECASE)
    if not match:
        return None
    start = max(match.start() - max_chars // 2, 0)
    end = min(match.end() + max_chars // 2, len(text))
    snippet = text[start:end]
    if start > 0:
        snippet = "…" + snippet
    if end < len(text):
        snippet += "…"
    def highlight_match(m): return f"<mark>{m.group(0)}</mark>"
    return re.sub(f"({re.escape(query)})", highlight_match, snippet, flags=re.IGNORECASE).strip()

# =========================
#   Sökning i Dokument (utan "Läs mer")
# =========================
def search_documents(query, visible_count=5):
    if not query or len(query.strip()) < 2:
        return "❗️ Skriv minst 2 tecken för att söka.", gr.update(visible=False)

    query = query.strip()
    results = []
    for doc in documents:
        score = 0
        filename_match = fuzz.partial_ratio(query.lower(), doc['filename'].lower()) > 80
        content_match = fuzz.partial_ratio(query.lower(), doc['content'].lower()) > 80
        if filename_match: score += 60
        if content_match:  score += 10
        if score > 0: results.append((doc, score, filename_match))
    results.sort(key=lambda x: (x[1], x[2]), reverse=True)

    html_output = ""
    shown = 0
    for doc, score, _ in results:
        if shown >= visible_count:
            break

        highlighted_filename = re.sub(f"({re.escape(query)})", r"<mark>\1</mark>", doc['filename'], flags=re.IGNORECASE)
        snippet = extract_context_snippet(doc["content"], query) or "⚠️ Ingen träfftext hittades."

        # Nedladdningslänk
        try:
            with open(doc["path"], "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
                download_link = f"📥 <a href='data:application/octet-stream;base64,{b64}' download='{doc['filename']}'>Ladda ner filen</a>"
        except Exception as e:
            download_link = f"❌ Kunde inte ladda filen: {html.escape(str(e))}"

        html_output += f"<h4>📄 {highlighted_filename}</h4>"
        html_output += f"<div class='preview' style='background:#f6f6f6;padding:10px;border-radius:5px;margin:10px 0;'>{snippet}</div>"
        html_output += f"<p>🔍 <b>Matchningspoäng:</b> {round(score, 1)}</p>"
        html_output += f"<p>{download_link}</p>"
        html_output += "<hr>"
        shown += 1

    show_more_visible = shown < len(results)
    return (html_output if html_output else "❌ Inga träffar hittades.", gr.update(visible=show_more_visible))

# =======================================================
#  Ladda & strukturera Bibliotek.docx
# =======================================================
def parse_word_sections(path):
    doc = Document(path)
    sections = []
    current = {"heading": "", "text": "", "images": []}
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current["heading"]:
                sections.append(current)
            current = {"heading": para.text, "text": "", "images": []}
        else:
            current["text"] += para.text + "\n"
        for run in para.runs:
            if run.element.xpath('.//a:blip'):
                blip = run.element.xpath('.//a:blip')[0]
                rEmbed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                image_part = doc.part.related_parts[rEmbed]
                img_b64 = base64.b64encode(image_part.blob).decode()
                current["images"].append(img_b64)
    if current["heading"]:
        sections.append(current)
    return sections

word_doc_path = os.path.join(BASE_DIR, "quickSearch", "Bibliotek.docx")
word_sections = []
if os.path.exists(word_doc_path):
    try:
        word_sections = parse_word_sections(word_doc_path)
    except Exception as e:
        print("⚠️ Kunde inte läsa Bibliotek.docx:", e)

def search_word_doc(query, visible_count=5):
    if not query or len(query.strip()) < 2:
        return "❗️ Skriv minst 2 tecken för att söka.", gr.update(visible=False)

    query = query.strip()
    results = []
    for section in word_sections:
        score = 0
        for w in query.lower().split():
            if fuzz.partial_ratio(w, section['heading'].lower()) > 80: score += 100
        for w in query.lower().split():
            if fuzz.partial_ratio(w, section['text'].lower()) > 80:    score += 10
        if score > 0: results.append((section, score))
    results.sort(key=lambda x: x[1], reverse=True)

    html_output = ""
    shown = 0
    formatter = HtmlFormatter(style="friendly", noclasses=True)

    for section, score in results:
        if shown >= visible_count:
            break

        highlighted_heading = re.sub(f"({re.escape(query)})", r"<mark>\1</mark>", section['heading'], flags=re.IGNORECASE)
        snippet = extract_context_snippet(section["text"], query) or "⚠️ Ingen träfftext hittades."
        highlighted_code = highlight(section['text'], PythonLexer(), formatter)
        images_html = "".join(f"<img src='data:image/png;base64,{img}' style='max-width:100%;'><br>" for img in section['images'])

        html_output += f"<h4>📑 {highlighted_heading}</h4>"
        html_output += f"""
        <details class='result'>
          <summary>
            <div class='preview'>{snippet}</div>
            <span class='toggle'>Läs mer</span>
          </summary>
          <div class='full'>{highlighted_code}{images_html}<p>🔍 <b>Matchningspoäng:</b> {round(score, 1)}</p></div>
        </details>
        <hr>
        """
        shown += 1

    show_more_visible = shown < len(results)
    return (html_output if html_output else "❌ Inga träffar hittades.", gr.update(visible=show_more_visible))

# =======================================================
#        Markdown-stöd för Developer.md
# =======================================================
def parse_markdown_sections(md_text):
    lines = md_text.splitlines()
    sections, current = [], {"heading": "", "text": ""}
    for line in lines:
        if line.lstrip().startswith("#"):
            if current["heading"] or current["text"].strip():
                sections.append(current)
            heading = line.lstrip().lstrip("#").strip()
            current = {"heading": heading, "text": ""}
        else:
            current["text"] += (line + "\n")
    if current["heading"] or current["text"].strip():
        sections.append(current)
    return sections or [{"heading": "Untitled", "text": md_text}]

def render_markdown_section_to_html(section_text):
    formatter = HtmlFormatter(style="friendly", noclasses=True)
    out, lines = [], section_text.splitlines()
    in_code, code_lang, code_buf = False, "", []

    def flush_code():
        if not code_buf: return ""
        code_str = "\n".join(code_buf)
        try:
            if code_lang:
                lexer = lexers.get_lexer_by_name(code_lang, stripall=False)
            else:
                if code_str.lstrip().startswith("{"):
                    lexer = JsonLexer()
                elif "SELECT" in code_str.upper():
                    lexer = SqlLexer()
                else:
                    lexer = TextLexer()
        except Exception:
            lexer = TextLexer()
        return highlight(code_str, lexer, formatter)

    for raw in lines:
        s = raw.strip()
        if s.startswith("```") and not in_code:
            in_code, code_lang, code_buf = True, s.strip("`").strip(), []
            if code_lang.startswith("```"): code_lang = code_lang[3:].strip()
            continue
        if s.startswith("```") and in_code:
            out.append(flush_code()); in_code, code_lang, code_buf = False, "", []; continue
        if in_code: code_buf.append(raw)
        else: out.append(html.escape(raw) + "<br>")
    if in_code: out.append(flush_code())
    return "".join(out)

developer_md_path = os.path.join(BASE_DIR, "quickSearch", "Developer.md")

def load_developer_sections_fresh():
    if not os.path.exists(developer_md_path): return []
    with open(developer_md_path, "r", encoding="utf-8") as f:
        return parse_markdown_sections(f.read())

def search_developer_doc(query, visible_count=5):
    if not query or len(query.strip()) < 2:
        return "❗️ Skriv minst 2 tecken för att söka.", gr.update(visible=False)
    results, sections = [], load_developer_sections_fresh()
    for sec in sections:
        score = 0
        for w in query.lower().split():
            if fuzz.partial_ratio(w, sec['heading'].lower()) > 80: score += 100
        for w in query.lower().split():
            if fuzz.partial_ratio(w, sec['text'].lower()) > 80:    score += 10
        if score > 0: results.append((sec, score))
    results.sort(key=lambda x: x[1], reverse=True)

    html_output = ""
    shown = 0
    for sec, score in results:
        if shown >= visible_count: break
        highlighted_heading = re.sub(f"({re.escape(query)})", r"<mark>\1</mark>", sec['heading'], flags=re.IGNORECASE)
        snippet = extract_context_snippet(sec["text"], query) or "⚠️ Ingen träfftext hittades."
        full_html = render_markdown_section_to_html(sec["text"])

        html_output += f"<h4>🛠️ {highlighted_heading}</h4>"
        html_output += f"""
        <details class='result'>
          <summary>
            <div class='preview'>{snippet}</div>
            <span class='toggle'>Läs mer</span>
          </summary>
          <div class='full'>{full_html}<p>🔍 <b>Matchningspoäng:</b> {round(score,1)}</p></div>
        </details>
        <hr>
        """
        shown += 1

    show_more_visible = shown < len(results)
    return (html_output if html_output else "❌ Inga träffar hittades.", gr.update(visible=show_more_visible))

# =========================
#          UI & CSS
# =========================
custom_css = """
/* Gemensam look för preview-rutor i Allmänt & Developer (med details.result) */
details.result > summary .preview {
  display: block;
  background:#f6f6f6;
  padding:10px;
  border-radius:5px;
  margin:10px 0;
}
/* För Allmänt & Developer – göm preview när details är öppet */
details.result[open] > summary .preview { display: none; }

/* Ikoner och klickbarhet för details */
details.result > summary {
  cursor: pointer;
  list-style: '▶️ ';
}
details.result[open] > summary {
  list-style: '▼ ';
}

span.toggle { font-weight: 600; }
"""

with gr.Blocks(css=custom_css) as demo:
    gr.Markdown("# 📚 NoWaste Dokumentbibliotek")

    # --- Dokument (ALLT synligt, ingen details) ---
    with gr.Tab("Dokument"):
        q1 = gr.Textbox(label="🔎 Sök i Dokument", placeholder="Ex: inventering, pall, artikelnummer")
        o1 = gr.HTML()
        more1 = gr.Button("⬇️ Visa fler", visible=False)
        vis1 = gr.State(5)

        def more_docs(q, v):
            html_out, show_more = search_documents(q, v + 5)
            return html_out, show_more, v + 5

        q1.change(fn=search_documents, inputs=[q1, vis1], outputs=[o1, more1])
        more1.click(fn=more_docs, inputs=[q1, vis1], outputs=[o1, more1, vis1])

    # --- Allmänt (Bibliotek.docx) ---
    with gr.Tab("Allmänt"):
        q2 = gr.Textbox(label="🔎 Sök i Bibliotek", placeholder="Ex: lager, capabilities, brandfarligt")
        o2 = gr.HTML()
        more2 = gr.Button("⬇️ Visa fler", visible=False)
        vis2 = gr.State(5)

        def more_word(q, v):
            html_out, show_more = search_word_doc(q, v + 5)
            return html_out, show_more, v + 5

        q2.change(fn=search_word_doc, inputs=[q2, vis2], outputs=[o2, more2])
        more2.click(fn=more_word, inputs=[q2, vis2], outputs=[o2, more2, vis2])

    # --- Developer (Developer.md) ---
    with gr.Tab("Developer"):
        q3 = gr.Textbox(label="🔎 Sök i Developer.md", placeholder="Ex: fraktsedlar, VM-Etikett, SELECT ...")
        o3 = gr.HTML()
        more3 = gr.Button("⬇️ Visa fler", visible=False)
        vis3 = gr.State(5)

        def more_dev(q, v):
            html_out, show_more = search_developer_doc(q, v + 5)
            return html_out, show_more, v + 5

        q3.change(fn=search_developer_doc, inputs=[q3, vis3], outputs=[o3, more3])
        more3.click(fn=more_dev, inputs=[q3, vis3], outputs=[o3, more3, vis3])

if __name__ == "__main__":
    demo.launch()
